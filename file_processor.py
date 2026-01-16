from fastapi import APIRouter, UploadFile, HTTPException, Depends, File
from typing import Dict, Any, Optional
import os
import json
import yaml
import pandas as pd
import PyPDF2
from docx import Document
import aiofiles
from pathlib import Path
import hashlib
import time

router = APIRouter(prefix="/files", tags=["File Processing"])

UPLOAD_DIR = "uploads/files"
os.makedirs(UPLOAD_DIR, exist_ok=True)

async def process_text_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo de texto simples"""
    async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
        content = await f.read()
    return {
        "type": "text",
        "content": content,
        "lines": len(content.split('\n')),
        "characters": len(content)
    }

async def process_pdf_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo PDF"""
    try:
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
        
        return {
            "type": "pdf",
            "content": text[:1000],  # Primeiros 1000 caracteres
            "pages": len(pdf_reader.pages),
            "full_text_length": len(text)
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao processar PDF: {str(e)}")

async def process_word_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo Word (.docx)"""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        
        return {
            "type": "word",
            "content": text[:1000],
            "paragraphs": len(doc.paragraphs),
            "full_text_length": len(text)
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao processar Word: {str(e)}")

async def process_csv_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo CSV"""
    try:
        df = pd.read_csv(file_path)
        return {
            "type": "csv",
            "rows": len(df),
            "columns": df.columns.tolist(),
            "preview": df.head(5).to_dict('records')
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao processar CSV: {str(e)}")

async def process_excel_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo Excel"""
    try:
        df = pd.read_excel(file_path)
        return {
            "type": "excel",
            "rows": len(df),
            "columns": df.columns.tolist(),
            "preview": df.head(5).to_dict('records')
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao processar Excel: {str(e)}")

async def process_json_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo JSON"""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
            data = json.loads(content)
        
        return {
            "type": "json",
            "content": data,
            "size": len(str(data))
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao processar JSON: {str(e)}")

async def process_xml_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo XML"""
    try:
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
            content = await f.read()
        
        return {
            "type": "xml",
            "content": content[:1000],
            "size": len(content)
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao processar XML: {str(e)}")

async def process_yaml_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo YAML"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
        
        return {
            "type": "yaml",
            "content": data,
            "size": len(str(data))
        }
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro ao processar YAML: {str(e)}")

async def process_parquet_file(file_path: str) -> Dict[str, Any]:
    """Processa arquivo Parquet - requer pyarrow"""
    raise HTTPException(
        status_code=400, 
        detail="Processamento de Parquet não disponível. Instale pyarrow: pip install pyarrow"
    )

@router.post("/process")
async def process_file(file: UploadFile = File(...)):
    """
    Processa qualquer tipo de arquivo suportado
    """
    try:
        # Salvar arquivo com timestamp único
        file_ext = file.filename.split('.')[-1].lower()
        timestamp = int(time.time())
        safe_filename = f"{timestamp}_{file.filename}"
        file_path = os.path.join(UPLOAD_DIR, safe_filename)
        
        content = await file.read()
        checksum = hashlib.sha256(content).hexdigest()
        
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)
        
        # Processar baseado na extensão
        if file_ext in ['txt', 'md', 'rtf']:
            result = await process_text_file(file_path)
        elif file_ext == 'pdf':
            result = await process_pdf_file(file_path)
        elif file_ext in ['doc', 'docx']:
            result = await process_word_file(file_path)
        elif file_ext == 'csv':
            result = await process_csv_file(file_path)
        elif file_ext in ['xlsx', 'xls', 'ods']:
            result = await process_excel_file(file_path)
        elif file_ext == 'json':
            result = await process_json_file(file_path)
        elif file_ext == 'xml':
            result = await process_xml_file(file_path)
        elif file_ext in ['yaml', 'yml']:
            result = await process_yaml_file(file_path)
        elif file_ext == 'parquet':
            result = await process_parquet_file(file_path)
        else:
            raise HTTPException(status_code=400, detail=f"Tipo de arquivo não suportado: {file_ext}")
        
        return {
            "success": True,
            "file_name": file.filename,
            "stored_name": safe_filename,
            "checksum": checksum,
            "size": len(content),
            **result
        }
    
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Erro ao processar arquivo: {str(e)}")
